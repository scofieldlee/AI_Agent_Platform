"""
Unified file parser for chat attachments.

Supports:
- PDF (.pdf)         → text extraction via pypdf
- Word (.docx)       → text extraction via python-docx
- Excel (.xlsx)      → table → text via openpyxl
- Text (.txt, .md)   → direct read
- Image (.jpg, .png) → base64 encode for multimodal LLM
- Video (.mp4, etc.) → metadata extraction (simple version)

Each parse returns a dict with:
  type:    "text" | "image" | "video"
  content: extracted text (for text type) | base64 string (for image type)
  meta:    { filename, mime_type, size_bytes, ... }
"""

import base64
import logging
import os
from pathlib import Path
from typing import Optional, Dict, Any, List

logger = logging.getLogger(__name__)

# --- File type mapping ---

EXT_TO_CATEGORY = {
    ".txt": "text",
    ".md": "text",
    ".csv": "text",
    ".pdf": "pdf",
    ".docx": "word",
    ".doc": "word",
    ".xlsx": "excel",
    ".xls": "excel",
    ".jpg": "image",
    ".jpeg": "image",
    ".png": "image",
    ".gif": "image",
    ".webp": "image",
    ".bmp": "image",
    ".mp4": "video",
    ".avi": "video",
    ".mov": "video",
    ".mkv": "video",
}

CATEGORY_TO_INPUT_TYPE = {
    "text": "text",
    "pdf": "pdf",
    "word": "word",
    "excel": "excel",
    "image": "image",
    "video": "video",
}

MIME_TYPES = {
    ".txt": "text/plain",
    ".md": "text/markdown",
    ".csv": "text/csv",
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".doc": "application/msword",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".xls": "application/vnd.ms-excel",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png": "image/png",
    ".gif": "image/gif",
    ".webp": "image/webp",
    ".bmp": "image/bmp",
    ".mp4": "video/mp4",
    ".avi": "video/x-msvideo",
    ".mov": "video/quicktime",
    ".mkv": "video/x-matroska",
}

MAX_TEXT_LENGTH = 50000  # Truncate extracted text to avoid context overflow


def get_file_category(filename: str) -> Optional[str]:
    """Get the category for a file based on its extension."""
    ext = Path(filename).suffix.lower()
    return EXT_TO_CATEGORY.get(ext)


def get_input_type(filename: str) -> Optional[str]:
    """Get the input_type identifier (for agent config validation)."""
    cat = get_file_category(filename)
    if cat:
        return CATEGORY_TO_INPUT_TYPE.get(cat)
    return None


def get_mime_type(filename: str) -> str:
    """Get MIME type for a file."""
    ext = Path(filename).suffix.lower()
    return MIME_TYPES.get(ext, "application/octet-stream")


async def parse_file(
    file_path: str,
    filename: str,
    content: Optional[bytes] = None,
) -> Dict[str, Any]:
    """Parse an uploaded file and return structured content.

    Args:
        file_path: Path to the saved file on disk.
        filename: Original filename (used to determine type).
        content: Optional raw bytes (if file is in memory, avoids re-read).

    Returns:
        Dict with keys:
          - type: "text" | "image" | "video"
          - content: str (text content or base64 for images)
          - meta: dict with filename, mime_type, size_bytes, and type-specific info
    """
    category = get_file_category(filename)
    if not category:
        return {
            "type": "text",
            "content": f"[不支持的文件类型: {filename}]",
            "meta": {"filename": filename, "mime_type": get_mime_type(filename), "size_bytes": 0, "error": "unsupported_type"},
        }

    mime = get_mime_type(filename)
    size = os.path.getsize(file_path) if os.path.exists(file_path) else (len(content) if content else 0)

    meta_base = {
        "filename": filename,
        "mime_type": mime,
        "size_bytes": size,
        "category": category,
    }

    try:
        if category == "text":
            text = _parse_text(file_path, content)
            return {"type": "text", "content": text, "meta": meta_base}

        elif category == "pdf":
            text = _parse_pdf(file_path, content)
            return {"type": "text", "content": text, "meta": meta_base}

        elif category == "word":
            text = _parse_docx(file_path, content)
            return {"type": "text", "content": text, "meta": meta_base}

        elif category == "excel":
            text = _parse_xlsx(file_path, content)
            return {"type": "text", "content": text, "meta": meta_base}

        elif category == "image":
            b64 = _parse_image(file_path, content)
            meta_base["base64"] = b64
            return {"type": "image", "content": b64, "meta": meta_base}

        elif category == "video":
            info = _parse_video(file_path, content)
            meta_base.update(info)
            return {"type": "video", "content": "", "meta": meta_base}

        else:
            return {"type": "text", "content": f"[无法解析的文件: {filename}]", "meta": meta_base}

    except Exception as e:
        logger.error(f"File parse failed for {filename}: {e}", exc_info=True)
        return {
            "type": "text",
            "content": f"[文件解析失败: {filename} — {str(e)[:100]}]",
            "meta": {**meta_base, "error": str(e)[:200]},
        }


# --- Individual parsers ---

def _read_bytes(file_path: str, content: Optional[bytes]) -> bytes:
    """Read file bytes from path or use provided content."""
    if content is not None:
        return content
    with open(file_path, "rb") as f:
        return f.read()


def _parse_text(file_path: str, content: Optional[bytes]) -> str:
    """Parse plain text / markdown / csv files."""
    raw = _read_bytes(file_path, content)
    text = raw.decode("utf-8", errors="replace")
    if len(text) > MAX_TEXT_LENGTH:
        text = text[:MAX_TEXT_LENGTH] + "\n\n[... 文本过长，已截断 ...]"
    return text


def _parse_pdf(file_path: str, content: Optional[bytes]) -> str:
    """Extract text from PDF using pypdf."""
    from pypdf import PdfReader

    raw = _read_bytes(file_path, content)
    import io
    reader = PdfReader(io.BytesIO(raw))

    parts: List[str] = []
    total_chars = 0
    for i, page in enumerate(reader.pages):
        page_text = page.extract_text() or ""
        if page_text.strip():
            parts.append(f"--- 第 {i + 1} 页 ---\n{page_text.strip()}")
            total_chars += len(page_text)
            if total_chars > MAX_TEXT_LENGTH:
                parts.append("\n[... PDF 内容过长，已截断 ...]")
                break

    return "\n\n".join(parts) if parts else "[PDF 文件无可提取的文本，可能是扫描件]"


def _parse_docx(file_path: str, content: Optional[bytes]) -> str:
    """Extract text from Word .docx using python-docx."""
    from docx import Document

    raw = _read_bytes(file_path, content)
    import io
    doc = Document(io.BytesIO(raw))

    parts: List[str] = []
    total_chars = 0

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
            total_chars += len(text)
            if total_chars > MAX_TEXT_LENGTH:
                parts.append("[... 文档内容过长，已截断 ...]")
                break

    # Extract tables
    if total_chars < MAX_TEXT_LENGTH:
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip(" |"):
                    parts.append(row_text)
                    total_chars += len(row_text)
                    if total_chars > MAX_TEXT_LENGTH:
                        parts.append("[... 文档内容过长，已截断 ...]")
                        break
            if total_chars > MAX_TEXT_LENGTH:
                break

    return "\n".join(parts) if parts else "[Word 文档无可提取的文本]"


def _parse_xlsx(file_path: str, content: Optional[bytes]) -> str:
    """Extract text from Excel .xlsx using openpyxl (sheet → table format)."""
    from openpyxl import load_workbook

    raw = _read_bytes(file_path, content)
    import io
    wb = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)

    parts: List[str] = []
    total_chars = 0

    for ws in wb.worksheets:
        sheet_lines: List[str] = [f"--- 工作表: {ws.title} ---"]
        for row in ws.iter_rows(values_only=True):
            row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
            if row_text.strip(" |"):
                sheet_lines.append(row_text)
                total_chars += len(row_text)
                if total_chars > MAX_TEXT_LENGTH:
                    sheet_lines.append("[... Excel 内容过长，已截断 ...]")
                    break
        parts.append("\n".join(sheet_lines))
        if total_chars > MAX_TEXT_LENGTH:
            break

    wb.close()
    return "\n\n".join(parts) if parts else "[Excel 文件无可提取的数据]"


def _parse_image(file_path: str, content: Optional[bytes]) -> str:
    """Read image and return base64 encoded string for multimodal LLM."""
    raw = _read_bytes(file_path, content)

    # Optionally resize large images to keep base64 manageable
    try:
        from PIL import Image
        import io as _io

        img = Image.open(_io.BytesIO(raw))
        # If image is very large, resize to max 1024px on longest side
        max_dim = 1024
        if max(img.size) > max_dim:
            ratio = max_dim / max(img.size)
            new_size = (int(img.size[0] * ratio), int(img.size[1] * ratio))
            img = img.resize(new_size, Image.LANCZOS)

        # Convert to RGB if needed (e.g., RGBA, P mode)
        if img.mode in ("RGBA", "P", "LA"):
            img = img.convert("RGB")

        buf = _io.BytesIO()
        img.save(buf, format="JPEG", quality=85)
        raw = buf.getvalue()
    except Exception as e:
        logger.warning(f"Image resize failed, using original: {e}")

    b64 = base64.b64encode(raw).decode("utf-8")
    return b64


def _parse_video(file_path: str, content: Optional[bytes]) -> Dict[str, Any]:
    """Extract video metadata (simple version — no frame extraction).

    Tries ffprobe first; falls back to basic file info.
    """
    import subprocess

    meta: Dict[str, Any] = {}

    # Try ffprobe for metadata
    try:
        result = subprocess.run(
            [
                "ffprobe", "-v", "quiet", "-print_format", "json",
                "-show_format", "-show_streams", file_path,
            ],
            capture_output=True, text=True, timeout=10,
        )
        if result.returncode == 0:
            import json
            data = json.loads(result.stdout)
            fmt = data.get("format", {})
            meta["duration_seconds"] = float(fmt.get("duration", 0))
            meta["format_name"] = fmt.get("format_name", "")

            # Find video stream for resolution
            for stream in data.get("streams", []):
                if stream.get("codec_type") == "video":
                    meta["width"] = stream.get("width", 0)
                    meta["height"] = stream.get("height", 0)
                    meta["video_codec"] = stream.get("codec_name", "")
                    break
    except FileNotFoundError:
        meta["note"] = "ffprobe not installed, basic metadata only"
    except Exception as e:
        meta["note"] = f"ffprobe failed: {str(e)[:100]}"

    return meta


def format_attachment_for_prompt(parsed: Dict[str, Any]) -> str:
    """Format a parsed attachment as text for injection into LLM prompt.

    Used for text-type attachments (pdf, word, excel, text).
    Image and video types are handled differently (base64 / metadata).
    """
    meta = parsed.get("meta", {})
    filename = meta.get("filename", "unknown")
    content = parsed.get("content", "")

    if parsed["type"] == "text":
        return f"[附件: {filename}]\n{content}"
    elif parsed["type"] == "video":
        duration = meta.get("duration_seconds", 0)
        w = meta.get("width", 0)
        h = meta.get("height", 0)
        return f"[视频附件: {filename}] (时长: {duration:.1f}s, 分辨率: {w}x{h}) — 视频内容暂不支持自动分析，请告知用户已收到视频并转人工处理。"
    else:
        return ""  # Images are handled via multimodal messages, not text
